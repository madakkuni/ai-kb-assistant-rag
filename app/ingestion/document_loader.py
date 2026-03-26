import os
import aiofiles
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pdfplumber
import docx
import markdown
from app.utils.logger import logger
from app.utils.config import settings

class DocumentIngestor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""]
        )

    async def ingest_file(self, file_path: str, filename: str) -> List[Document]:
        logger.info(f"Ingesting file: {filename}")
        try:
            ext = os.path.splitext(filename)[1].lower()
            text = ""
            
            if ext == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            elif ext == '.docx':
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif ext in ['.txt', '.md']:
                async with aiofiles.open(file_path, mode='r', encoding='utf-8') as f:
                    text = await f.read()
            else:
                raise ValueError(f"Unsupported file format: {ext}")
            
            # Create a single document and split
            base_doc = Document(page_content=text, metadata={"source": filename})
            docs = self.text_splitter.split_documents([base_doc])
            
            # Add page/chunk numbers to metadata
            for i, doc in enumerate(docs):
                doc.metadata["chunk"] = i
                
            logger.info(f"Successfully chunked {filename} into {len(docs)} chunks.")
            return docs
            
        except Exception as e:
            logger.error(f"Error ingesting file {filename}: {str(e)}", exc_info=True)
            raise
